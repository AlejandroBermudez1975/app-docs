import streamlit as st
import pandas as pd
import os
import re
from io import BytesIO
from PyPDF2 import PdfReader
import docx

# Función para extraer texto de archivos PDF
def extract_text_from_pdf(file):
    text = ""
    try:
        reader = PdfReader(file)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text
    except:
        pass
    return text if text.strip() else "[Sin texto detectable]"

# Función para extraer texto de archivos Word (incluye tablas)
def extract_text_from_word(file):
    text = ""
    try:
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except:
        pass
    return text if text.strip() else "[Sin texto detectable]"

# Función mejorada para extraer emails
def extract_email(text):
    pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    match = re.search(pattern, text)
    return match.group(0) if match else ""

# Interfaz de la app
st.set_page_config(page_title="Extractor de Emails", layout="centered")

st.title("📄 Extraer Emails de Documentos")

uploaded_files = st.file_uploader("Sube hasta 100 archivos PDF o Word", type=["pdf", "docx"], accept_multiple_files=True)

if uploaded_files:
    progress_text = "⏳ Procesando archivos..."
    progress_bar = st.progress(0, text=progress_text)

    data = []

    for i, file in enumerate(uploaded_files):
        filename = file.name
        ext = os.path.splitext(filename)[1].lower()
        text = ""

        if ext == ".pdf":
            text = extract_text_from_pdf(file)
        elif ext == ".docx":
            text = extract_text_from_word(file)

        email = extract_email(text)

        data.append({
            "email": email,
            "archivo": filename
        })

        progress_bar.progress((i + 1) / len(uploaded_files), text=f"📁 Procesando: {filename}")

    df = pd.DataFrame(data)
    df = df[["email", "archivo"]]  # Solo columnas necesarias

    output = BytesIO()
    df.to_excel(output, index=False, engine="openpyxl")
    st.success("✅ ¡Extracción completada!")

    st.download_button(
        label="📥 Descargar Excel",
        data=output.getvalue(),
        file_name="resultados.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
